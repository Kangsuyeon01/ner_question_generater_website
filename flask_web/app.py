# app.py
import sys
import os
sys.path.append('/home/suyeon/code/capstone2/bert_ner')
from flask import Flask, render_template, request,send_file
from question_generater import start
from time import time,sleep
import json
from docx import Document

#Flask 객체 인스턴스 생성
app = Flask(__name__)

@app.route('/') # 접속하는 url
def index():

  return render_template('index.html')

@app.route('/qg_main', methods = ['POST', 'GET'])
def content_crawl():
    global questions 
    if request.method == 'POST':
        context =  request.form["input_context"]
        empty_questions,empty_answers, ox_questions, ox_answers = start(context)
        questions = [empty_questions,empty_answers, ox_questions, ox_answers]
    
    return render_template("qg_main.html", questions = questions)

@app.route('/ox_quiz', methods = ['POST', 'GET'])
def ox_quiz():

  return render_template("ox_quiz.html", questions = questions)

@app.route('/qg_print', methods = ['POST', 'GET'])
def download():
    path = question_print()
    return send_file(path, as_attachment=True)

@app.route('/qg_result', methods = ['POST', 'GET'])
def qg_result():
  
  return render_template("qg_result.html", questions = questions)


def make_jsonfile(questions):
  global unix_time
  unix_time = int(time())
  QA_data = {"empty_qa": {"questions":questions[0],"answers":questions[1]} ,"ox_qa":{"questions":questions[2],"answers":questions[3]}}
  file_path = "./flask_web/QA_data/" + str(unix_time) + ".json"
  with open(file_path, 'w') as outfile:
      json.dump(QA_data, outfile, indent=4)
  return file_path

def load_jsondile(file_path):
  with open(file_path, 'r') as f:
    data = json.load(f)
  return data


def question_print():
  global questions
  # 워드 문서 생성하기
  doc = Document()
  
  # 문단 추가하기 - 빈칸
  doc.add_paragraph('빈칸 문제')
  for i in range(len(questions[0])):
    doc.add_paragraph("Q"+str(i+1)+". " + questions[0][i])

  # 페이지 추가하기
  doc.add_page_break()
  # 문단 추가하기 - OX
  doc.add_paragraph('OX 문제')
  for i in range(len(questions[2])):
    doc.add_paragraph("Q"+str(i+1)+". " + questions[2][i])

  # 페이지 추가하기
  doc.add_page_break()
  # 문단 추가하기 - 빈칸 답
  doc.add_paragraph('빈칸 문제 답')
  for i in range(len(questions[1])):
    tmp = ", ".join(questions[1][i])
    doc.add_paragraph("A"+str(i+1)+". " + tmp)

  # 페이지 추가하기
  doc.add_page_break()
  # 문단 추가하기 - OX 답
  doc.add_paragraph('OX 문제 답')
  for i in range(len(questions[3])):
    tmp = ", ".join(questions[3][i])
    doc.add_paragraph("A"+str(i+1)+". " + tmp )

  unix_time = int(time())
  file_path = "/home/suyeon/code/capstone2/flask_web/docs/" +str(unix_time) + '.docx'
  doc.save(file_path)
  
  return file_path
 
  

if __name__=="__main__":
  # app.run(debug=True)
  # host 등을 직접 지정하고 싶다면
  app.run(host="127.0.0.1", port="5555", debug=True)

  # flask run --host=0.0.0.0 로 서버 실행